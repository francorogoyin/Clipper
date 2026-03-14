"""
Exportador a DOCX usando python-docx.

"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from Backend.App.Exportadores.Exportador_Base import (
    Exportador_Base,
)

ALINEACIONES = {
    "izquierda": WD_ALIGN_PARAGRAPH.LEFT,
    "derecha": WD_ALIGN_PARAGRAPH.RIGHT,
    "centro": WD_ALIGN_PARAGRAPH.CENTER,
    "justificada": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class Exportador_Docx(Exportador_Base):

    """
    Genera un archivo DOCX con los highlights.

    """

    @property
    def Extension(self) -> str:
        return "docx"

    def Exportar(
        self,
        Dataframe: pd.DataFrame,
        Ruta_Salida: str,
        Config_Estilo: dict,
    ) -> str:

        """
        Genera el DOCX.

        """

        Doc = Document()

        for _, Fila in Dataframe.iterrows():
            Texto = str(Fila.get("Texto", ""))
            Autor = str(
                Fila.get("Autor", "") or ""
            )
            Libro = str(
                Fila.get("Libro", "") or ""
            )

            Campos_Sub = Config_Estilo.get(
                "Sub_Line_Campos",
                ["autor", "libro"],
            )
            Partes = []
            if "autor" in Campos_Sub and Autor:
                Partes.append(Autor)
            if "libro" in Campos_Sub and Libro:
                Partes.append(Libro)
            Sub_Line = " — ".join(Partes)

            # Highlight.
            P_H = Doc.add_paragraph()
            Run_H = P_H.add_run(Texto)
            Run_H.font.size = Pt(
                Config_Estilo.get(
                    "Highlight_Tamano", 12
                )
            )
            Color_H = Config_Estilo.get(
                "Highlight_Color", "#000000"
            ).lstrip("#")
            Run_H.font.color.rgb = RGBColor(
                int(Color_H[0:2], 16),
                int(Color_H[2:4], 16),
                int(Color_H[4:6], 16),
            )
            Formato_H = Config_Estilo.get(
                "Highlight_Formato", "normal"
            )
            if "negrita" in Formato_H:
                Run_H.bold = True
            if "cursiva" in Formato_H:
                Run_H.italic = True
            P_H.alignment = ALINEACIONES.get(
                Config_Estilo.get(
                    "Highlight_Alineacion",
                    "izquierda",
                ),
                WD_ALIGN_PARAGRAPH.LEFT,
            )

            # Sub_Line.
            if Sub_Line:
                P_S = Doc.add_paragraph()
                Run_S = P_S.add_run(Sub_Line)
                Run_S.font.size = Pt(
                    Config_Estilo.get(
                        "Sub_Line_Tamano", 10
                    )
                )
                Color_S = Config_Estilo.get(
                    "Sub_Line_Color", "#666666"
                ).lstrip("#")
                Run_S.font.color.rgb = RGBColor(
                    int(Color_S[0:2], 16),
                    int(Color_S[2:4], 16),
                    int(Color_S[4:6], 16),
                )
                Formato_S = Config_Estilo.get(
                    "Sub_Line_Formato", "cursiva"
                )
                if "negrita" in Formato_S:
                    Run_S.bold = True
                if "cursiva" in Formato_S:
                    Run_S.italic = True

            # Separador.
            Doc.add_paragraph()

        Doc.save(Ruta_Salida)
        return Ruta_Salida
